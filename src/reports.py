from docx import Document
from docx.shared import Inches
from datetime import datetime

def cust_form_letter(template_path, output_path, data):
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if '{{' + key + '}}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{{' + key + '}}', str(value))

    # Update placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if '{{' + key + '}}' in cell.text:
                        cell.text = cell.text.replace('{{' + key + '}}', str(value))

    # Saving the populated document
    doc.save(output_path)

current_date = datetime.now()
formatted_date = current_date.strftime("%B %d, %Y")

# Example data (you would replace these with actual calculated values)
data = {
    'date': formatted_date,
    'customer_name': 'Alpha Packaging',
    'address': '123 Main St',
    'city': 'Cleveland',
    'state': 'OH',
    'zip_code': '44134',
    'fault_loc': 'Secondary',
    'tr_kva': '300',
    'pri_volt': '13.2',
    'sec_volt': '120/208',
    'no_phases': '3',
    'tr_conn': 'WYE-WYE solidly grounded',
    'tr_imp': '4.9',
    'x_r_ratio': '4.5',
    'fault_voltage_lvl': '208',
    '3ph_fault': '29765',
    'l_g_fault': '27654',
    'x_r_pos_ratio': '4.53',
    'x_r_zero_ratio': '3.71',
    'pro_device': '50E Fuse',
    'engineer_sig': 'James T. Long, PE',
}

cust_form_letter(r'C:\Users\51192\OneDrive - FirstEnergy Corp\Documents\Projects\Customer Fault Calcs\Short Circuit Form Letter Template.docx', r'C:\Users\51192\OneDrive - FirstEnergy Corp\Documents\Projects\Customer Fault Calcs\Short Circuit Form Letter Test.docx', data)
